"""UDPE Word 模板导入器: docx → grid schemaJson.

设计文档: plans/udpe-design/outputs/m3-stage4-design.md §四 3.2

核心能力:
- 解析段落 (含占位符 {{ xxx | filter }})
- 标题样式 (Heading 1/2/3) → title 类型
- 表格 → grid 类型 (每行对应 grid row, 每个单元格对应 cell)
- 样式映射 (bold / fontSize / color / align)
- 多表格 + 段落混合布局

V1 限制:
- 不支持复杂表格合并 (Word 表格合并靠 merge_fields)
- 不支持图片、图表、形状
- 不支持页眉页脚
- 不支持目录、脚注
"""
import io
import logging
import re
from typing import Any, Optional

logger = logging.getLogger(__name__)

# 占位符正则: {{ path | filter1 | filter2 }}
PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")


def _clean_template_text(text: str) -> str:
    """规范化占位符: {{x}} → {{ x }}, {{x|filter}} → {{ x | filter }}."""
    def repl(m: re.Match) -> str:
        inner = m.group(1).strip()
        return "{{ " + inner + " }}"
    return PLACEHOLDER_RE.sub(repl, text)


def _extract_paragraph_data(paragraph) -> dict:
    """提取段落数据 (文本、样式、对齐、字体)."""
    text = paragraph.text.strip()
    if not text:
        return {}
    
    # 样式映射
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    is_heading = "heading" in style_name
    
    # 对齐方式
    align = None
    if paragraph.alignment:
        align_map = {
            0: "left",    # WD_ALIGN_PARAGRAPH.LEFT
            1: "center",  # WD_ALIGN_PARAGRAPH.CENTER
            2: "right",   # WD_ALIGN_PARAGRAPH.RIGHT
            3: "justify", # WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        align = align_map.get(paragraph.alignment, None)
    
    # 字体信息 (取第一个 run 的字体)
    bold = False
    font_size = None
    color = None
    if paragraph.runs:
        first_run = paragraph.runs[0]
        if first_run.font:
            bold = bool(first_run.font.bold)
            if first_run.font.size:
                # 转换为 pt (Word 用 EMU, 1 pt = 12700 EMU)
                font_size = int(first_run.font.size / 12700)
            if first_run.font.color and first_run.font.color.rgb:
                color = f"#{first_run.font.color.rgb}"
    
    return {
        "text": _clean_template_text(text),
        "bold": bold,
        **({"fontSize": font_size} if font_size else {}),
        **({"align": align} if align else {}),
        **({"color": color} if color else {}),
    }


def _convert_table(table) -> dict:
    """Word 表格 → grid block schemaJson."""
    if not table.rows:
        return {
            "type": "grid",
            "colCount": 1,
            "border": True,
            "rows": [],
        }
    
    # 计算列数 (取第一行的列数)
    col_count = len(table.rows[0].cells) if table.rows else 1
    
    # 构造 grid rows
    grid_rows = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            # 合并单元格处理: 计算 col_span
            # Word 表格合并比较复杂, V1 简化处理: 忽略合并, 每个 cell 独立
            cell_text = cell.text.strip()
            if not cell_text:
                cell_text = ""
            
            # 提取字体信息
            bold = False
            font_size = None
            color = None
            align = None
            
            if cell.paragraphs:
                first_para = cell.paragraphs[0]
                if first_para.runs:
                    first_run = first_para.runs[0]
                    if first_run.font:
                        bold = bool(first_run.font.bold)
                        if first_run.font.size:
                            font_size = int(first_run.font.size / 12700)
                        if first_run.font.color and first_run.font.color.rgb:
                            color = f"#{first_run.font.color.rgb}"
                if first_para.alignment:
                    align_map = {
                        0: "left",
                        1: "center",
                        2: "right",
                        3: "justify",
                    }
                    align = align_map.get(first_para.alignment, None)
            
            cell_data = {
                "text": _clean_template_text(cell_text),
                "span": 1,  # V1 不支持合并
                "bold": bold,
            }
            if font_size:
                cell_data["fontSize"] = font_size
            if align:
                cell_data["align"] = align
            if color:
                cell_data["color"] = color
            
            row_cells.append(cell_data)
        
        if row_cells:
            # 计算行高 (默认 14mm)
            height_mm = 14
            grid_rows.append({"height": height_mm, "cells": row_cells})
    
    return {
        "type": "grid",
        "colCount": col_count,
        "border": True,
        "borderColor": "#000000",
        "borderWidth": 1,
        "rows": grid_rows,
    }


def _convert_paragraph(paragraph) -> Optional[dict]:
    """段落 → title/text block."""
    data = _extract_paragraph_data(paragraph)
    if not data:
        return None
    
    text = data.get("text", "")
    if not text:
        return None
    
    # 占位符检查
    has_placeholder = bool(PLACEHOLDER_RE.search(text))
    
    # 样式映射
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    is_heading = "heading" in style_name
    
    if is_heading:
        # 标题样式 → title 类型
        # 提取标题级别 (Heading 1 → 1, Heading 2 → 2)
        level = 1
        if "heading 2" in style_name:
            level = 2
        elif "heading 3" in style_name:
            level = 3
        
        return {
            "type": "title",
            "text": text,
            "fontSize": data.get("fontSize", 22 if level == 1 else 18 if level == 2 else 16),
            "align": data.get("align", "center" if level == 1 else "left"),
            **({"bold": data["bold"]} if "bold" in data else {}),
            **({"color": data["color"]} if "color" in data else {}),
        }
    else:
        # 普通段落 → text 类型
        return {
            "type": "text",
            "text": text,
            "fontSize": data.get("fontSize", 11),
            "align": data.get("align", "left"),
            **({"bold": data["bold"]} if "bold" in data else {}),
            **({"color": data["color"]} if "color" in data else {}),
        }


def _convert_document(doc) -> dict:
    """Word 文档 → schemaJson (混合段落和表格)."""
    body_blocks = []
    
    # 遍历文档元素 (段落和表格)
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag == 'p':  # 段落
            # 通过 element 查找对应的 paragraph 对象
            # python-docx 的 API 需要从 document 对象获取
            # 这里我们直接遍历 doc.paragraphs 和 doc.tables
            pass
        elif tag == 'tbl':  # 表格
            pass
    
    # 更简单的方法: 按顺序遍历段落和表格
    # python-docx 的 doc.element.body 包含所有元素, 但顺序可能不准确
    # 我们改用 doc.paragraphs 和 doc.tables, 但需要保持顺序
    
    # 收集所有元素及其位置
    elements = []
    
    # 遍历段落
    for para in doc.paragraphs:
        # 获取段落在文档中的位置 (通过 element 对象)
        elements.append(('paragraph', para, para._element))
    
    # 遍历表格
    for table in doc.tables:
        elements.append(('table', table, table._element))
    
    # 按文档中的顺序排序 (通过 element 对象的 .getparent() 索引)
    # 这是一个简化方法, 实际上需要更复杂的排序
    # 我们改用 doc.element.body 的顺序
    
    # 重新实现: 直接遍历 doc.element.body
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'p':  # 段落
            # 查找对应的 paragraph 对象
            for para in doc.paragraphs:
                if para._element is child:
                    block = _convert_paragraph(para)
                    if block:
                        body_blocks.append(block)
                    break
        elif tag == 'tbl':  # 表格
            # 查找对应的 table 对象
            for table in doc.tables:
                if table._element is child:
                    grid = _convert_table(table)
                    if grid:
                        body_blocks.append(grid)
                    break
    
    return {"body": body_blocks}


def parse_docx(file_bytes: bytes) -> dict:
    """主入口: docx 字节流 → 解析文档结构.

    Returns:
        {
            "filename": "example.docx",
            "totalElements": 10,  # 段落 + 表格总数
            "schemaJson": { "body": [...] },
            "placeholders": ["form.title", "form.totalAmount", ...],
            "warnings": ["表格合并不支持, 已降级为独立单元格", ...],
        }
    """
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    
    # 转换文档
    schema_json = _convert_document(doc)
    
    # 提取占位符
    placeholders = set()
    warnings = []
    
    def extract_placeholders(obj):
        if isinstance(obj, dict):
            for key, value in obj.items():
                if key == "text" and isinstance(value, str):
                    for m in PLACEHOLDER_RE.finditer(value):
                        path = m.group(1).split("|")[0].strip()
                        if path:
                            placeholders.add(path)
                elif isinstance(value, (dict, list)):
                    extract_placeholders(value)
        elif isinstance(obj, list):
            for item in obj:
                extract_placeholders(item)
    
    extract_placeholders(schema_json)
    
    # 检查表格合并
    for block in schema_json.get("body", []):
        if block.get("type") == "grid":
            for row in block.get("rows", []):
                for cell in row.get("cells", []):
                    if cell.get("span", 1) > 1:
                        warnings.append("表格合并不支持, 已降级为独立单元格")
                        break
                if warnings:
                    break
            if warnings:
                break
    
    # 计算元素总数
    total_elements = 0
    for block in schema_json.get("body", []):
        if block.get("type") == "grid":
            # 表格算一个元素
            total_elements += 1
        else:
            # 段落算一个元素
            total_elements += 1
    
    return {
        "totalElements": total_elements,
        "schemaJson": schema_json,
        "placeholders": sorted(placeholders),
        "warnings": warnings,
    }
